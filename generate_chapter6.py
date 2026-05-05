#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Генератор Главы 6 диплома TariffRadar в формате .docx
Требует: pip install python-docx
"""

import sys
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("❌ Ошибка: не установлена библиотека python-docx")
    print("Установите её: pip install python-docx")
    sys.exit(1)

def set_margins(doc, left=30, right=10, top=20, bottom=20):
    """Устанавливает поля в миллиметрах"""
    sections = doc.sections
    for section in sections:
        section.left_margin = Inches(left / 25.4)
        section.right_margin = Inches(right / 25.4)
        section.top_margin = Inches(top / 25.4)
        section.bottom_margin = Inches(bottom / 25.4)

def set_paragraph_format(paragraph, space_before=0, space_after=0, line_height=18, indent_first=720):
    """Устанавливает формат абзаца согласно БГУИР СТП 01-2024"""
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    fmt.line_spacing_rule = 4  # Точно (atLeast)
    fmt.line_spacing = line_height
    fmt.first_line_indent = Inches(indent_first / 914400)

def add_chapter_title(doc, text):
    """Добавляет заголовок главы по центру, заглавными буквами"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    set_paragraph_format(para, space_before=0, space_after=200, indent_first=0)
    return para

def add_subsection(doc, text):
    """Добавляет подзаголовок подраздела (жирный, с отступа)"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0)
    para.paragraph_format.first_line_indent = Inches(1.25 / 2.54)  # 1,25 см
    run = para.add_run(text)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    set_paragraph_format(para, space_before=200, space_after=0, indent_first=0)
    return para

def add_text(doc, text):
    """Добавляет обычный текст параграфа"""
    para = doc.add_paragraph(text)
    for run in para.runs:
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
    set_paragraph_format(para, space_before=0, space_after=0)
    return para

def main():
    doc = Document()
    set_margins(doc)

    # Заголовок
    add_chapter_title(doc, '6 ТЕХНИКО-ЭКОНОМИЧЕСКОЕ ОБОСНОВАНИЕ И МОДЕЛЬ МОНЕТИЗАЦИИ')

    # Введение
    add_text(doc, 'Платформа TariffRadar разработана как двусторонний B2B SaaS маркетплейс, '
             'обслуживающий логистический рынок Беларуси. Ключевой задачей технико-экономического '
             'анализа является обоснование жизнеспособности бизнес-модели, оценка доходов по предложенным '
             'тарифам подписок, и определение путей масштабирования платформы.')

    # 6.1
    add_subsection(doc, '6.1 Прогноз доходов по тарифам подписок')

    add_text(doc, 'Экономическая модель платформы основана на SaaS-модели с тремя уровнями подписки. '
             'Каждый уровень предназначен для различных сегментов пользователей и обеспечивает разный '
             'набор функций, соответствующий потребностям и бюджету целевой аудитории.')

    # Таблица 6.1
    table1 = doc.add_table(rows=4, cols=4)
    table1.style = 'Light Grid'

    headers = ['Параметр', 'Free (Базовый)', 'Pro ($29/мес)', 'Business ($99/мес)']
    for i, h in enumerate(headers):
        table1.rows[0].cells[i].text = h
        for para in table1.rows[0].cells[i].paragraphs:
            for run in para.runs:
                run.font.bold = True

    data = [
        ('Поиск грузов и предложений', 'До 5 поисков/день', 'Неограниченно', 'Неограниченно'),
        ('История сделок и аналитика', '30 дней', '90 дней', '12 месяцев'),
        ('Экспорт CSV, API, Team', 'Нет', 'CSV только', 'Полный доступ'),
    ]

    for row_idx, row_data in enumerate(data, 1):
        for col_idx, text in enumerate(row_data):
            table1.rows[row_idx].cells[col_idx].text = text

    caption1 = doc.add_paragraph()
    caption1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption1.add_run('Таблица 6.1 – Функции по уровням подписки')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    set_paragraph_format(caption1, space_before=100, space_after=200, indent_first=0)

    add_text(doc, 'Free-тариф предоставляет базовый доступ к маркетплейсу, достаточный для пользователей, '
             'которые редко проводят анализ рынка. Pro-тариф ориентирован на активных пользователей маркетплейса, '
             'предоставляя неограниченный доступ к поиску. Business-тариф предназначен для крупных логистических '
             'операторов и корпоративных грузоотправителей.')

    add_text(doc, 'Прогноз доходов платформы основан на консервативной модели приобретения пользователей и конверсии. '
             'На первоначальном этапе предполагается следующее распределение базы пользователей:')

    # Таблица 6.2
    table2 = doc.add_table(rows=5, cols=3)
    table2.style = 'Light Grid'

    headers2 = ['Месяц', 'Пользователи', 'Ожидаемый MRR (USD)']
    for i, h in enumerate(headers2):
        table2.rows[0].cells[i].text = h
        for para in table2.rows[0].cells[i].paragraphs:
            for run in para.runs:
                run.font.bold = True

    forecast = [
        ('Месяц 1-3', '50 (40 Free / 8 Pro / 2 Bus)', '$300'),
        ('Месяц 4-6', '150 (110 Free / 32 Pro / 8 Bus)', '$1,100'),
        ('Месяц 7-12', '400 (280 Free / 90 Pro / 30 Bus)', '$3,600'),
        ('Итого в год', '~600 активных пользователей', '~$28,800'),
    ]

    for row_idx, (m, u, r) in enumerate(forecast, 1):
        table2.rows[row_idx].cells[0].text = m
        table2.rows[row_idx].cells[1].text = u
        table2.rows[row_idx].cells[2].text = r

    caption2 = doc.add_paragraph()
    caption2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption2.add_run('Таблица 6.2 – Прогноз доходов (MRR)')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    set_paragraph_format(caption2, space_before=100, space_after=200, indent_first=0)

    add_text(doc, 'Расчеты основаны на следующих предположениях: коэффициент конверсии Free→Pro — 20%, '
             'Free→Business — 4%; средняя стоимость Pro — $29, Business — $99; коэффициент удержания — 5% для Free, '
             '2% для Pro/Business. Прогноз консервативен.')

    # 6.2
    add_subsection(doc, '6.2 Анализ окупаемости и перспективы масштабирования')

    add_text(doc, 'Анализ окупаемости показывает, что платформа достигает безубыточности примерно на 9–12-й месяц. '
             'Основные категории расходов:')

    # Список
    expenses = [
        'Инфраструктура (хостинг, БД, кэш) — $800–1 200/месяц',
        'Персонал разработки — $3 000–5 000/месяц',
        'Платежные системы (Stripe) — 2,9% + $0,30 от транзакции',
        'Маркетинг и привлечение пользователей — $500–1 000/месяц',
        'Прочие (лицензии, инструменты) — $200–300/месяц',
    ]

    for expense in expenses:
        para = doc.add_paragraph(expense, style='List Bullet')
        for run in para.runs:
            run.font.size = Pt(14)
            run.font.name = 'Times New Roman'

    add_text(doc, 'Общие ежемесячные расходы на первоначальном этапе — $5 500–7 500. '
             'При увеличении базы до 2 000 пользователей (месяц 18–24), доход вырастет до ~$15 000–20 000, '
             'что обеспечит окупаемость инвестиций.')

    add_text(doc, 'Перспективы масштабирования определяются: '
             '(1) географическим расширением на рынки РФ, УК, КЗ; '
             '(2) углублением функциональности (GPS, ML-прогнозирование); '
             '(3) открытым API для интеграции с TMS; '
             '(4) монетизацией агрегированных данных; '
             '(5) мобильными приложениями.')

    add_text(doc, 'Платформа может обслужить до 10 000 пользователей к концу 3-го года, '
             'генерируя месячный доход $100 000–150 000. Это сделает TariffRadar привлекательным объектом '
             'для инвестиций и приобретения крупными логистическими операторами.')

    # Сохранение
    output = r'D:\диплом\Глава_6_Экономика.docx'
    doc.save(output)
    print(f'✅ Глава 6 создана успешно: {output}')

if __name__ == '__main__':
    main()
